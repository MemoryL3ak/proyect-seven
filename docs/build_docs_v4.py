"""Genera los DOCX de la documentación v4 a partir de sus markdown.

Basado en el motor de build_manual_v2_docx.py, parametrizado para varios
documentos y con tabla de contenidos autogenerada desde los encabezados.

Uso:  python build_docs_v4.py
Salida: MANUAL_USUARIO_v4.docx y FUNCIONALIDADES_PLATAFORMA_v11.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).parent

NAVY = RGBColor(0x0E, 0x2A, 0x47)
PRIMARY = RGBColor(0x1F, 0x4E, 0x8C)
ACCENT = RGBColor(0xE3, 0xA8, 0x08)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5E, 0x6B, 0x7A)
LIGHT_BG = RGBColor(0xF4, 0xF6, 0xFA)
TIP_BG = RGBColor(0xE7, 0xF5, 0xEC)
TIP_BORDER = RGBColor(0x2E, 0x7D, 0x32)
WARN_BG = RGBColor(0xFF, 0xF4, 0xD6)
WARN_BORDER = RGBColor(0xC7, 0x8C, 0x00)
CRIT_BG = RGBColor(0xFD, 0xE2, 0xE2)
CRIT_BORDER = RGBColor(0xB3, 0x1B, 0x1B)
NOTE_BG = RGBColor(0xE6, 0xEE, 0xF7)
NOTE_BORDER = RGBColor(0x1F, 0x4E, 0x8C)
TABLE_HEADER_BG = RGBColor(0x1F, 0x4E, 0x8C)
TABLE_ROW_ALT = RGBColor(0xF7, 0xF9, 0xFC)
CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")

CALLOUT_PATTERNS = [
    ("🟢", "TIP", TIP_BG, TIP_BORDER),
    ("🟡", "ATENCIÓN", WARN_BG, WARN_BORDER),
    ("🔴", "CRÍTICO", CRIT_BG, CRIT_BORDER),
    ("📌", "NOTA", NOTE_BG, NOTE_BORDER),
    ("⚠️", "ATENCIÓN", WARN_BG, WARN_BORDER),
]


def _rgb_hex(color):
    return "{:02X}{:02X}{:02X}".format(color[0], color[1], color[2])


def build_doc(cfg):
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    # ── XML helpers ────────────────────────────────────────────────────────
    def shade_cell(cell, color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), _rgb_hex(color))
        tc_pr.append(shd)

    def set_cell_borders(cell, color=RGBColor(0xD0, 0xD7, 0xDE), size="6"):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), size)
            b.set(qn("w:color"), _rgb_hex(color))
            borders.append(b)
        tc_pr.append(borders)

    def paragraph_background(paragraph, color):
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), _rgb_hex(color))
        p_pr.append(shd)

    def add_page_break():
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    def add_horizontal_rule(color=PRIMARY, width_pt=12):
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(width_pt))
        bottom.set(qn("w:color"), _rgb_hex(color))
        bottom.set(qn("w:space"), "1")
        pBdr.append(bottom)
        p_pr.append(pBdr)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)

    # ── Footer con número de página ────────────────────────────────────────
    footer = doc.sections[0].footer
    foot_p = footer.paragraphs[0]
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = foot_p.add_run(f"{cfg['footer']}   ·   Página ")
    run_title.font.size = Pt(8)
    run_title.font.color.rgb = MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = foot_p.add_run()
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = MUTED
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_sep)
    page_run._r.append(fld_end)

    # ── Portada ────────────────────────────────────────────────────────────
    band = doc.add_paragraph()
    band.paragraph_format.space_before = Pt(0)
    band.paragraph_format.space_after = Pt(0)
    paragraph_background(band, NAVY)
    r = band.add_run(" ")
    r.font.size = Pt(36)
    r.font.color.rgb = NAVY

    for _ in range(3):
        doc.add_paragraph()

    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = badge.add_run("SEVEN ARENA")
    br.bold = True
    br.font.size = Pt(14)
    br.font.color.rgb = ACCENT

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(cfg["title"])
    tr.bold = True
    tr.font.size = Pt(38)
    tr.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(cfg["subtitle"])
    sr.font.size = Pt(14)
    sr.font.color.rgb = MUTED
    sr.italic = True

    doc.add_paragraph()
    add_horizontal_rule(PRIMARY, 16)

    meta_tbl = doc.add_table(rows=len(cfg["meta"]), cols=2)
    meta_tbl.autofit = False
    for i, (k, v) in enumerate(cfg["meta"]):
        c1 = meta_tbl.cell(i, 0)
        c2 = meta_tbl.cell(i, 1)
        c1.width = Cm(4.5)
        c2.width = Cm(11)
        c1p = c1.paragraphs[0]
        c1p.paragraph_format.space_after = Pt(2)
        c1r = c1p.add_run(k.upper())
        c1r.bold = True
        c1r.font.size = Pt(9)
        c1r.font.color.rgb = MUTED
        c2p = c2.paragraphs[0]
        c2p.paragraph_format.space_after = Pt(2)
        c2r = c2p.add_run(v)
        c2r.font.size = Pt(11)
        c2r.font.color.rgb = INK
        for c in (c1, c2):
            shade_cell(c, LIGHT_BG if i % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF))
            set_cell_borders(c, RGBColor(0xE1, 0xE5, 0xEB))
    add_page_break()

    # ── Contenido markdown ─────────────────────────────────────────────────
    md = Path(cfg["md"]).read_text(encoding="utf-8")
    lines = md.split("\n")

    # TOC autogenerada desde los encabezados (# PARTE/ANEXOS y ##)
    toc_title = doc.add_paragraph()
    tr = toc_title.add_run("Tabla de Contenidos")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = NAVY
    add_horizontal_rule(ACCENT, 10)
    in_code_toc = False
    for line in lines:
        s = line.rstrip()
        if s.strip().startswith("```"):
            in_code_toc = not in_code_toc
            continue
        if in_code_toc:
            continue
        is_part = s.startswith("# ") and ("PARTE" in s or "ANEXOS" in s)
        is_section = s.startswith("## ") and not s.startswith("## Acerca")
        if not (is_part or is_section):
            continue
        text = s[2:].strip() if is_part else "  " + s[3:].strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        if is_part:
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = PRIMARY
            p.paragraph_format.space_before = Pt(6)
        else:
            r.font.size = Pt(10.5)
            r.font.color.rgb = INK
    add_page_break()

    # ── Renderers ──────────────────────────────────────────────────────────
    def add_inline(paragraph, text, base_size=11, base_color=INK):
        for part in INLINE_RE.split(text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = paragraph.add_run(part[2:-2])
                r.bold = True
                r.font.size = Pt(base_size)
                r.font.color.rgb = base_color
            elif part.startswith("`") and part.endswith("`"):
                r = paragraph.add_run(part[1:-1])
                r.font.name = "Consolas"
                r.font.size = Pt(base_size - 1)
                r.font.color.rgb = RGBColor(0xB3, 0x1B, 0x1B)
            else:
                r = paragraph.add_run(part)
                r.font.size = Pt(base_size)
                r.font.color.rgb = base_color

    def add_part_header(text):
        add_page_break()
        banner = doc.add_paragraph()
        banner.paragraph_format.space_before = Pt(6)
        banner.paragraph_format.space_after = Pt(0)
        paragraph_background(banner, NAVY)
        run = banner.add_run(" " + text + " ")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        bar = doc.add_paragraph()
        bar.paragraph_format.space_before = Pt(0)
        bar.paragraph_format.space_after = Pt(14)
        paragraph_background(bar, ACCENT)
        r = bar.add_run(" ")
        r.font.size = Pt(2)

    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = PRIMARY
        add_horizontal_rule(ACCENT, 8)

    def add_subsection_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = NAVY

    def add_sub_subsection_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11.5)
        r.font.color.rgb = PRIMARY

    def try_render_callout(text):
        for emoji, label, bg, border in CALLOUT_PATTERNS:
            if text.startswith(emoji):
                body = text[len(emoji):].strip()
                body = re.sub(r"^\*\*[^*]+\*\*\s*:?\s*", "", body)
                tbl = doc.add_table(rows=1, cols=1)
                tbl.autofit = True
                cell = tbl.cell(0, 0)
                shade_cell(cell, bg)
                tc_pr = cell._tc.get_or_add_tcPr()
                borders = OxmlElement("w:tcBorders")
                for side in ("top", "bottom", "right"):
                    b = OxmlElement(f"w:{side}")
                    b.set(qn("w:val"), "single")
                    b.set(qn("w:sz"), "4")
                    b.set(qn("w:color"), _rgb_hex(border))
                    borders.append(b)
                left = OxmlElement("w:left")
                left.set(qn("w:val"), "single")
                left.set(qn("w:sz"), "32")
                left.set(qn("w:color"), _rgb_hex(border))
                borders.append(left)
                tc_pr.append(borders)
                pl = cell.paragraphs[0]
                pl.paragraph_format.space_after = Pt(2)
                lr = pl.add_run(f"{emoji}  {label}")
                lr.bold = True
                lr.font.size = Pt(10)
                lr.font.color.rgb = border
                pb = cell.add_paragraph()
                pb.paragraph_format.space_after = Pt(0)
                add_inline(pb, body, base_size=10.5, base_color=INK)
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(4)
                return True
        return False

    def parse_markdown_table(all_lines, start):
        rows = []
        i = start
        while i < len(all_lines) and all_lines[i].strip().startswith("|"):
            row = [c.strip() for c in all_lines[i].strip().strip("|").split("|")]
            rows.append(row)
            i += 1
        if len(rows) >= 2 and all(set(c) <= set("-: ") for c in rows[1]):
            rows.pop(1)
        return rows, i

    def render_table(rows):
        if not rows:
            return
        n_cols = len(rows[0])
        tbl = doc.add_table(rows=len(rows), cols=n_cols)
        tbl.autofit = True
        tbl.style = "Table Grid"
        for j, header in enumerate(rows[0]):
            cell = tbl.cell(0, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade_cell(cell, TABLE_HEADER_BG)
            set_cell_borders(cell, TABLE_HEADER_BG, size="4")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(header.replace("**", ""))
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for i, row in enumerate(rows[1:], start=1):
            for j, value in enumerate(row):
                if j >= n_cols:
                    break
                cell = tbl.cell(i, j)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if i % 2 == 0:
                    shade_cell(cell, TABLE_ROW_ALT)
                set_cell_borders(cell, RGBColor(0xDD, 0xE3, 0xEA))
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                add_inline(p, value, base_size=10.5, base_color=INK)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def render_code_block(code_lines):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = True
        cell = tbl.cell(0, 0)
        shade_cell(cell, CODE_BG)
        set_cell_borders(cell, RGBColor(0xD0, 0xD7, 0xDE))
        first = True
        for line in code_lines:
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(line if line else " ")
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = NAVY
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def add_bullet(text, level=0):
        style_name = "List Bullet" if level == 0 else "List Bullet 2"
        try:
            p = doc.add_paragraph(style=style_name)
        except KeyError:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75 * (level + 1))
            p.add_run("• ").bold = True
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, text)

    def add_numbered(text):
        try:
            p = doc.add_paragraph(style="List Number")
        except KeyError:
            p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, text)

    def add_paragraph_text(text):
        if text.startswith("*") and text.endswith("*") and not text.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text.strip("*").strip())
            r.italic = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = MUTED
            return
        if try_render_callout(text):
            return
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, text)

    # ── Recorrido del markdown ─────────────────────────────────────────────
    i = 0
    first_h1_skipped = False
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if line.startswith("# "):
            text = line[2:].strip()
            if "PARTE" in text or "ANEXOS" in text:
                add_part_header(text)
            elif not first_h1_skipped:
                first_h1_skipped = True  # título del doc: ya está en la portada
            else:
                add_part_header(text)
            i += 1
            continue

        if line.startswith("## "):
            add_section_header(line[3:].strip())
            i += 1
            continue
        if line.startswith("### "):
            add_subsection_header(line[4:].strip())
            i += 1
            continue
        if line.startswith("#### "):
            add_sub_subsection_header(line[5:].strip())
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            render_code_block(code_lines)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, end = parse_markdown_table(lines, i)
            render_table(rows)
            i = end
            continue

        if stripped.startswith("> "):
            body = stripped[2:].strip()
            if not try_render_callout(body):
                try_render_callout("📌 " + body)
            i += 1
            continue

        if stripped == "---" or not stripped:
            i += 1
            continue

        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_num:
            add_numbered(m_num.group(2).strip())
            i += 1
            continue

        if stripped.startswith("- "):
            level = 1 if raw.startswith(("  - ", "   - ", "    - ")) else 0
            add_bullet(stripped[2:].strip(), level=level)
            i += 1
            continue

        add_paragraph_text(stripped)
        i += 1

    doc.save(cfg["out"])
    print(f"Saved: {cfg['out']}")


CONFIGS = [
    {
        "md": HERE / "MANUAL_USUARIO_v4.md",
        "out": HERE / "MANUAL_USUARIO_v4.docx",
        "title": "Manual de Usuario",
        "subtitle": "Plataforma integral para la gestión de eventos deportivos",
        "footer": "Manual de Usuario · Seven Arena",
        "meta": [
            ("Versión", "4.0"),
            ("Fecha", "Agosto 2026"),
            ("Dirigido a", "Comité Organizador, Operaciones, Hotelería, Transporte, "
                            "Alimentación, Acreditación, Prensa, Voluntariado, Directores "
                            "de Delegación y personal técnico."),
        ],
    },
    {
        "md": HERE / "FUNCIONALIDADES_PLATAFORMA_v11.md",
        "out": HERE / "FUNCIONALIDADES_PLATAFORMA_v11.docx",
        "title": "Funcionalidades de la Plataforma",
        "subtitle": "Descripción funcional de los módulos operativos de Seven Arena",
        "footer": "Funcionalidades de la Plataforma · Seven Arena",
        "meta": [
            ("Versión", "4.0 (v11)"),
            ("Fecha", "Agosto 2026"),
            ("Alcance", "Módulos operativos, portales de usuario final, aplicación móvil "
                         "nativa y funcionalidades liberadas entre julio y agosto de 2026."),
        ],
    },
]

if __name__ == "__main__":
    for cfg in CONFIGS:
        build_doc(cfg)
