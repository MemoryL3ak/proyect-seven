"""Convert REACT_NATIVE_HANDOFF.md to REACT_NATIVE_HANDOFF.docx"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
MD = (HERE / "REACT_NATIVE_HANDOFF.md").read_text(encoding="utf-8")
OUT = HERE / "REACT_NATIVE_HANDOFF_v2.docx"

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

TEAL = RGBColor(0x21, 0xD0, 0xB3)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GREY = RGBColor(0x64, 0x74, 0x8B)

def add_inline(paragraph, text: str):
    """Parse inline markdown (**bold**, `code`, [text](link))"""
    parts = re.split(r'(\*\*.+?\*\*|`[^`]+`|\[.+?\]\(.+?\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif part.startswith("[") and "](" in part:
            m = re.match(r'\[(.+?)\]\((.+?)\)', part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = TEAL
                run.underline = True
        else:
            paragraph.add_run(part)

def add_heading(text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = TEAL
        run.bold = True
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = DARK
        run.bold = True
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = DARK
        run.bold = True
    else:
        run.font.size = Pt(11)
        run.bold = True

def add_code_block(code: str, lang: str = ""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code.rstrip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = DARK

def add_table_from_md(rows):
    if not rows:
        return
    cols = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            tcell = tbl.cell(i, j)
            p = tcell.paragraphs[0]
            add_inline(p, cell.strip())
            if i == 0:
                for run in p.runs:
                    run.bold = True

# Parse markdown
lines = MD.split("\n")
i = 0
in_code = False
code_buf = []
code_lang = ""
in_table = False
table_buf = []

while i < len(lines):
    line = lines[i]

    # Code blocks
    if line.strip().startswith("```"):
        if not in_code:
            in_code = True
            code_lang = line.strip()[3:]
            code_buf = []
        else:
            add_code_block("\n".join(code_buf), code_lang)
            in_code = False
        i += 1
        continue
    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # Tables
    if "|" in line and line.strip().startswith("|"):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Skip separator row (---|---)
        if all(re.match(r'^-+:?$|^:?-+:?$', c) for c in cells):
            i += 1
            continue
        table_buf.append(cells)
        in_table = True
        i += 1
        continue
    elif in_table:
        add_table_from_md(table_buf)
        doc.add_paragraph()
        table_buf = []
        in_table = False

    # Headings
    m = re.match(r'^(#{1,4})\s+(.+)$', line)
    if m:
        add_heading(m.group(2).strip(), len(m.group(1)))
        i += 1
        continue

    # Horizontal rule
    if re.match(r'^-{3,}$', line.strip()):
        doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # Blockquote
    if line.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line[2:])
        run.italic = True
        run.font.color.rgb = GREY
        i += 1
        continue

    # Unordered list
    if re.match(r'^[-*]\s+', line):
        text = re.sub(r'^[-*]\s+', '', line)
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, text)
        i += 1
        continue

    # Ordered list
    if re.match(r'^\d+\.\s+', line):
        text = re.sub(r'^\d+\.\s+', '', line)
        p = doc.add_paragraph(style="List Number")
        add_inline(p, text)
        i += 1
        continue

    # Empty line
    if not line.strip():
        i += 1
        continue

    # Regular paragraph
    p = doc.add_paragraph()
    add_inline(p, line)
    i += 1

# Flush any remaining table
if in_table and table_buf:
    add_table_from_md(table_buf)

doc.save(OUT)
print(f"Saved: {OUT}")
