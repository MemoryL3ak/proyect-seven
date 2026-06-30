# -*- coding: utf-8 -*-
"""
Genera docs/FUNCIONALIDADES_PLATAFORMA_v10.docx a partir de la estructura
definida en este script. Es una actualización del v9 con los módulos
implementados recientemente.

Uso:
    python docs/generate_funcionalidades_v10.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────────────────────────────────────────
# CONTENIDO
# ──────────────────────────────────────────────────────────────────────────────
# Cada entrada del documento es una tupla (tipo, payload).
#   ("TITLE", str)            — Título principal del documento
#   ("SUBTITLE", str)         — Subtítulo bajo el título
#   ("H1", str)               — Título de sección (1., 2., ...)
#   ("INTRO", str)            — Párrafo introductorio (cursiva opcional)
#   ("H2", str)               — Subtítulo "a) Funcionalidad" / "b) Módulo..."
#   ("UL", list[str])         — Lista con viñetas
#   ("P", str)                — Párrafo simple
#
# Las **negritas** en los textos se procesan automáticamente con marcadores **.

DOC = [
    ("TITLE", "Funcionalidades de la Plataforma Seven Arena"),

    # ─────────────────── 1. Participantes ───────────────────
    ("H1", "1. Gestión de Participantes y Delegaciones"),
    ("INTRO", "Permite mantener un registro centralizado de todos los participantes del evento (atletas, jefes de delegación, VIPs, prensa, oficiales) organizados por delegación o país, con la información personal, deportiva y logística necesaria para la operación."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Registro individual y masivo de participantes asociados a una delegación, indicando nombre, documento de identidad, nacionalidad, género, fecha de nacimiento, deporte, rol y tipo de cliente (TA, VIP, Prensa, Oficial).",
        "Clasificación de participantes mediante atributos operativos: jefe de delegación, requiere acreditación, requiere transporte, requiere hospedaje, requiere alimentación.",
        "Asociación de participantes a hospedajes, vuelos de llegada y salida, y vehículos asignados durante la estadía.",
        "Visualización consolidada de la delegación con totales por categoría, estado de acreditación y servicios contratados.",
        "Edición de los datos del participante en cualquier momento de la operación, manteniendo trazabilidad de los cambios realizados.",
        "Exportación del listado de participantes a archivos en formato XLSX para reportería externa.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Formulario de registro individual con validación de campos obligatorios y verificación de duplicados por documento de identidad.",
        "Carga masiva de participantes mediante archivo XLSX con plantilla descargable, validación previa de cada fila y reporte de errores antes de la confirmación de la carga.",
        "Edición masiva de atributos comunes (delegación, tipo de cliente, requerimientos) sobre selección múltiple.",
        "Eliminación lógica de registros con confirmación explícita por parte del usuario administrador.",
    ]),

    # ─────────────────── 2. Proveedores ───────────────────
    ("H1", "2. Gestión de Proveedores y Recursos Operativos"),
    ("INTRO", "Permite registrar a los proveedores externos contratados para la prestación de servicios (transporte, alimentación, hospedaje, seguridad, etc.) y a los recursos humanos y materiales que cada proveedor pone a disposición del evento."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Registro de proveedores con razón social, RUT, contacto, tipo de servicio prestado y datos bancarios.",
        "Asociación de participantes del proveedor (conductores, personal de cocina, personal de seguridad) con identificación del rol específico mediante banderas operativas (ej. esChofer = sí).",
        "Registro de vehículos del proveedor de transporte con tipo, patente, capacidad máxima, modelo y año.",
        "Carga de fotografías del personal del proveedor para fines de identificación visual y acreditación.",
        "Carga y visualización de documentos asociados al personal (licencia de conducir, antecedentes, certificaciones), con almacenamiento seguro y versionado.",
        "Edición individual y masiva de la información del proveedor y de sus recursos.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Formulario de alta de proveedor con validación de RUT y verificación de unicidad.",
        "Carga individual y masiva de participantes del proveedor mediante archivo XLSX.",
        "Carga de fotografía del participante en formato JPG o PNG, con previsualización antes de la confirmación.",
        "Carga de documentos en formato PDF o imagen, con metadato del tipo de documento y fecha de vencimiento.",
        "Eliminación lógica de proveedores y participantes con confirmación explícita.",
    ]),

    # ─────────────────── 3. Calendario ───────────────────
    ("H1", "3. Calendario Deportivo y Operacional"),
    ("INTRO", "Permite construir y consultar el calendario completo del evento, integrando competencias deportivas, traslados, ceremonias, entrenamientos y actividades complementarias, con vista diaria y por disciplina."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Creación de eventos deportivos con fecha, hora, sede, disciplina, fase, género y delegaciones participantes.",
        "Visualización del calendario en formato mensual, semanal y diario, con filtros por disciplina, sede y delegación.",
        "Vista por día consolidada con todas las actividades programadas (competencias, traslados, ceremonias, entrenamientos).",
        "Asociación automática entre eventos deportivos y solicitudes de transporte generadas para los participantes involucrados.",
        "Edición de eventos con notificación automática a los actores afectados (delegaciones, conductores, transporte).",
        "Exportación del calendario a archivos en formato XLSX y PDF para distribución externa.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Formulario de creación de evento con autocompletado de sedes y disciplinas previamente registradas.",
        "Carga masiva de eventos mediante archivo XLSX con validación de fechas, sedes y disciplinas.",
        "Edición individual y masiva de eventos sobre selección múltiple desde la vista de listado.",
        "Eliminación de eventos con confirmación explícita y notificación a los actores asociados.",
    ]),

    # ─────────────────── 4. Transporte ───────────────────
    ("H1", "4. Gestión de Transporte"),
    ("INTRO", "Permite administrar de forma integral la operación de transporte del evento, desde la solicitud del servicio hasta la confirmación de la finalización del viaje, incluyendo la asignación de vehículos, conductores, tarifas y trazabilidad en tiempo real."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Recepción de solicitudes de transporte generadas desde los portales de usuario (VIP, jefe de delegación, atleta) o creación manual desde el módulo de administración.",
        "Clasificación de viajes por tipo: solo ida, ida y regreso, viaje desde portal, traslado operativo.",
        "Asignación de vehículo y conductor a cada viaje, con filtrado automático del listado de conductores disponibles según la capacidad solicitada y el tipo de vehículo registrado.",
        "Cálculo automático de la tarifa del viaje según tipo de vehículo, distancia, tipo de viaje (ida o ida y regreso) y horario.",
        "Estados del viaje: solicitado, programado, en curso, recogido, finalizado, cancelado.",
        "Cancelación de viajes con registro automático del usuario y la fecha de la cancelación en la bitácora del viaje.",
        "Bitácora del viaje (registro histórico) con trazabilidad de todas las acciones ejecutadas: creación, asignación de conductor, cambio de vehículo, inicio de ruta, recogida del pasajero, finalización y observaciones manuales.",
        "Visualización de los viajes históricos con filtros por estado, fecha, conductor, delegación y tipo de viaje.",
        "Reportería operativa con totales de viajes por conductor, vehículo, delegación y rango de fechas.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Formulario de creación manual de viaje con autocompletado de direcciones mediante integración con servicios de mapas.",
        "Edición de viajes en cualquier estado anterior a la finalización, con registro automático de cada cambio en la bitácora.",
        "Carga masiva de viajes mediante archivo XLSX con validación de fechas, sedes, conductores y vehículos.",
        "Eliminación lógica de viajes con confirmación explícita y registro del usuario que ejecutó la acción.",
    ]),

    # ─────────────────── 5. Portal Conductor ───────────────────
    ("H1", "5. Portal del Conductor"),
    ("INTRO", "Aplicación destinada a los conductores asociados a proveedores de transporte, mediante la cual reciben sus viajes asignados, ejecutan las acciones operativas del viaje y reportan su posición en tiempo real al centro de control."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Visualización de la bandeja de viajes asignados al conductor, con diferenciación visual entre viajes vistos y no vistos (similar a una bandeja de correo electrónico).",
        "Acciones operativas sobre cada viaje: iniciar ruta, marcar pasajero recogido (con confirmación mediante código de pasajero), marcar viaje finalizado.",
        "Envío de la posición geográfica del conductor en tiempo real al centro de control durante la ejecución del viaje.",
        "Carga de fotografías del vehículo y del pasajero al momento de la recogida y la finalización del viaje.",
        "Carga y visualización de los documentos personales del conductor (licencia, antecedentes) directamente desde el portal.",
        "Solicitud automática de permisos de ubicación cuando el conductor ejecuta una acción que los requiere.",
        "Mantención de la sesión activa y del envío de posición incluso cuando el dispositivo se encuentra en segundo plano (mediante mecanismos de wake lock y aplicación nativa).",
        "Comunicación bidireccional con el centro de control mediante chat embebido en cada viaje.",
        "Visualización de la credencial digital con código QR dentro de la propia aplicación, sin necesidad de abrir una ventana adicional, con opción de imprimirla desde la misma vista.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Acceso al portal mediante credenciales generadas por el administrador, vinculadas al participante del proveedor con bandera esChofer = sí.",
        "Persistencia local del estado de los viajes vistos y no vistos mediante almacenamiento del dispositivo.",
        "Registro automático de cada acción ejecutada en la bitácora del viaje.",
    ]),

    # ─────────────────── 6. Portal VIP ───────────────────
    ("H1", "6. Portal de Solicitud de Vehículo (VIP)"),
    ("INTRO", "Aplicación destinada a participantes registrados como tipo de cliente VIP, mediante la cual pueden solicitar de forma autónoma los servicios de transporte que requieran durante su estadía. El portal incorpora además las secciones operativas de premiaciones, cupones y alimentación, garantizando paridad funcional con los portales del jefe de delegación y del atleta."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Generación de solicitudes de transporte indicando origen, destino, fecha, hora, cantidad de pasajeros y tipo de vehículo requerido.",
        "Selección del tipo de vehículo entre las opciones disponibles (Sedán, SUV, Van, Minibus, Bus) con visualización de la capacidad máxima asociada.",
        "Validación automática de la coherencia entre la cantidad de pasajeros solicitada y la capacidad del vehículo seleccionado.",
        "Visualización del listado de solicitudes propias con su estado actual (solicitado, programado, en curso, finalizado, cancelado).",
        "Modificación o cancelación de solicitudes propias mientras el viaje no haya iniciado.",
        "Recepción de notificaciones sobre el estado del viaje y la asignación del conductor.",
        "Sección de premiaciones con indicadores agregados (programadas, realizadas, confirmadas, pendientes), vista de calendario mensual con marcadores en los días que tienen ceremonias y vista de lista intercambiables, filtros por estado, rol y asistencia, y opción de confirmar o declinar la asistencia a cada ceremonia.",
        "Sección de cupones con dos vistas, disponibles y mis cupones, catálogo de beneficios del evento, reclamo del cupón y visualización del código QR con código de respaldo para el canje en el comercio.",
        "Sección de alimentación con el menú de hoy, el menú de mañana y los lugares de comida del evento, sin filtros por tipo de cliente.",
        "Visualización de la credencial digital dentro de la propia aplicación, sin necesidad de abrir una ventana adicional, con opción de imprimirla desde la misma vista.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Acceso al portal mediante credenciales generadas por el administrador, vinculadas al participante registrado con tipo de cliente VIP.",
        "Formulario de solicitud con autocompletado de direcciones mediante integración con servicios de mapas.",
        "Confirmación visual del envío de la solicitud y del cambio de estado del viaje.",
    ]),

    # ─────────────────── 7. Portal Jefe de Delegación ───────────────────
    ("H1", "7. Portal del Jefe de Delegación"),
    ("INTRO", "Aplicación destinada a los participantes registrados como tipo de cliente TA con la bandera de jefe de delegación activa, mediante la cual administran las solicitudes y consultas de su delegación durante el evento."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Visualización del calendario de actividades deportivas de su delegación.",
        "Generación de solicitudes de transporte para los miembros de la delegación.",
        "Consulta de la información de hospedaje, alimentación y sedes asignadas.",
        "Visualización de la información de la cuenta y datos de la delegación.",
        "Acceso a las secciones operativas del portal: itinerario, actividades, calendario, sedes, alimentación, premiaciones, cupones, delegación y cuenta.",
        "Sección de premiaciones con vista de calendario mensual y vista de lista intercambiables, filtros por estado, disciplina, sede y búsqueda libre, agrupación por día y resumen de entregadores por rol (oro, plata, bronce, autoridad y premiador).",
        "Sección de cupones con catálogo completo y reclamo del código QR sin requerir inicio de sesión adicional.",
        "Visualización de la alimentación sin filtros por tipo de cliente: todos los miembros de la delegación pueden consultar la oferta completa de menús y lugares de comida.",
        "Visualización de la credencial digital dentro de la propia aplicación, sin necesidad de abrir una ventana adicional.",
        "Recepción de notificaciones operativas relacionadas con la delegación.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Acceso al portal mediante credenciales generadas por el administrador, vinculadas al participante con tipo de cliente TA y bandera jefe de delegación = sí.",
        "Filtrado automático de las secciones visibles del portal según el rol del usuario.",
    ]),

    # ─────────────────── 8. Portal Atleta ───────────────────
    ("H1", "8. Portal del Atleta"),
    ("INTRO", "Aplicación destinada a los participantes registrados como tipo de cliente TA sin la bandera de jefe de delegación, mediante la cual consultan la información operativa propia durante el evento."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Consulta del calendario personal de actividades deportivas.",
        "Consulta de la información de hospedaje, alimentación y sedes asignadas.",
        "Visualización de la información de la cuenta y datos personales.",
        "Acceso a las secciones operativas del portal: actividades, calendario, sedes, alimentación, premiaciones, cupones y cuenta.",
        "Sección de premiaciones con vista de calendario y vista de lista intercambiables, además de filtros equivalentes a los del portal del jefe de delegación.",
        "Sección de cupones con catálogo y reclamo del código QR.",
        "Visualización de la alimentación sin filtros por tipo de cliente: catálogo completo de menús y lugares de comida.",
        "Visualización de la credencial digital dentro de la propia aplicación, sin abrir una ventana adicional.",
        "Recepción de notificaciones operativas relacionadas con sus actividades.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Acceso al portal mediante credenciales generadas por el administrador, vinculadas al participante con tipo de cliente TA y bandera jefe de delegación = no.",
        "Filtrado automático de las secciones visibles del portal según el rol del usuario.",
    ]),

    # ─────────────────── 9. Tracking GPS ───────────────────
    ("H1", "9. Tracking GPS y Mapa Operativo"),
    ("INTRO", "Permite visualizar en tiempo real la ubicación de todos los conductores activos del evento, monitorear los viajes en curso y reaccionar de forma inmediata frente a desviaciones o incidentes."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Mapa central con la posición en tiempo real de todos los conductores activos, actualizada cada pocos segundos.",
        "Diferenciación visual de los conductores según estado del viaje (en ruta al origen, en ruta al destino, sin viaje activo).",
        "Visualización del recorrido histórico de un viaje específico con la trazabilidad completa de las posiciones reportadas.",
        "Filtrado del mapa por proveedor de transporte, conductor, tipo de vehículo y estado del viaje.",
        "Alertas automáticas frente a desviaciones de ruta, detenciones prolongadas o pérdida de señal.",
        "Pantalla de monitoreo de conductores con indicadores agregados (conductores en línea, posiciones GPS recientes, viajes activos), tabla detallada del estado de cada conductor y mapa con sus posiciones en tiempo real, diferenciando visualmente a los conductores conectados de los desconectados y mostrando el detalle de cada uno al seleccionarlo.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Registro automático de cada posición reportada por los portales del conductor en la base de datos histórica.",
        "Persistencia del recorrido completo de cada viaje para fines de auditoría y reportería.",
    ]),

    # ─────────────────── 10. Heatmap ───────────────────
    ("H1", "10. Panel de Conductores y Heatmap"),
    ("INTRO", "Permite visualizar la distribución horaria de la carga de trabajo de los conductores durante un día operativo, facilitando la planificación y la detección de desbalances en la asignación de viajes."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Matriz visual de conductores en filas y horas en columnas (de 07:00 a 23:00) con intensidad de color proporcional a la cantidad de viajes asignados en cada franja.",
        "Indicadores agregados (KPIs) con el total de viajes del día, conductores activos, vehículos en operación y horas totales conducidas.",
        "Ranking de conductores por cantidad de viajes ejecutados.",
        "Semáforos de carga (subutilizado, óptimo, sobrecargado) por conductor.",
        "Tooltip al pasar sobre cada celda con el detalle de los viajes de esa franja horaria.",
        "Selector de fecha para consultar días anteriores y actualización automática cada quince segundos durante el día en curso.",
        "Integración con la pantalla de monitoreo de conductores descrita en la sección 9, permitiendo visualizar en una sola vista la carga histórica de trabajo y la presencia GPS de los conductores en tiempo real.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Generación automática del heatmap a partir de la información de viajes registrados en el módulo de transporte.",
    ]),

    # ─────────────────── 11. Tarifas ───────────────────
    ("H1", "11. Gestión de Tarifas"),
    ("INTRO", "Permite mantener el listado de tarifas vigentes para cada tipo de vehículo y tipo de viaje, asegurando el cálculo correcto del valor del servicio en cada solicitud generada."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Registro de tarifas por tipo de vehículo (Sedán, SUV, Van, Minibus, Bus) y tipo de viaje (ida, ida y regreso).",
        "Vigencia de tarifas con fecha de inicio y fecha de término.",
        "Aplicación automática de la tarifa correspondiente al momento de generar un viaje.",
        "Visualización del histórico de tarifas para fines de auditoría.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Formulario de creación de tarifa con validación de unicidad por combinación de vehículo, tipo de viaje y vigencia.",
        "Edición y eliminación de tarifas con confirmación explícita.",
    ]),

    # ─────────────────── 12. Vuelos ───────────────────
    ("H1", "12. Tracking de Vuelos"),
    ("INTRO", "Permite registrar los vuelos de llegada y salida de los participantes y monitorear su estado en tiempo real mediante integración con servicios externos de información aeronáutica."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Registro del número de vuelo, aerolínea, fecha de llegada o salida, origen y destino para cada participante.",
        "Consulta automática del estado del vuelo (programado, en ruta, aterrizado, retrasado, cancelado) mediante integración con la API de AviationStack.",
        "Visualización del horario estimado de llegada actualizado en tiempo real.",
        "Generación automática de alertas frente a retrasos significativos para coordinar la operación de transporte de recibimiento.",
        "Diferenciación entre vuelos pasados, en curso y futuros, priorizando la información del vuelo activo o más próximo.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Formulario de registro de vuelo con validación del número de vuelo y la fecha contra la información retornada por la API externa.",
        "Carga masiva de vuelos mediante archivo XLSX con validación previa de los datos.",
    ]),

    # ─────────────────── 13. Notificaciones ───────────────────
    ("H1", "13. Sistema de Notificaciones"),
    ("INTRO", "Permite mantener informados a los usuarios de la plataforma sobre eventos relevantes para su operación mediante notificaciones in-app y correo electrónico."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Notificaciones in-app accesibles mediante la campana del menú principal, con indicador del número de notificaciones no leídas.",
        "Diferenciación visual de las notificaciones según el tipo (transporte, calendario, confirmación, mensaje, alerta) mediante íconos representativos.",
        "Marcado de notificaciones como leídas individualmente o de forma masiva.",
        "Envío opcional de notificaciones por correo electrónico para eventos críticos.",
        "Generación automática de notificaciones en respuesta a cambios de estado en los módulos de transporte, calendario y administración de usuarios.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Persistencia de todas las notificaciones generadas en la base de datos, asociadas al usuario destinatario.",
        "Registro de la fecha y hora de envío y de la fecha y hora de lectura para cada notificación.",
    ]),

    # ─────────────────── 14. Sofia IA ───────────────────
    ("H1", "14. Asistente de Inteligencia Artificial (SofIA)"),
    ("INTRO", "Permite a los usuarios administradores realizar consultas operativas en lenguaje natural sobre el estado de la plataforma y obtener respuestas inmediatas con información actualizada del sistema."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Interfaz de chat embebida en la plataforma de administración, accesible desde cualquier vista.",
        "Procesamiento de consultas en lenguaje natural sobre viajes, conductores, vehículos, participantes, tarifas y eventos del calendario.",
        "Acceso del asistente a las funciones internas de la plataforma mediante herramientas integradas, permitiendo respuestas con datos en tiempo real.",
        "Historial de conversaciones por usuario para retomar consultas previas.",
        "Respuestas con formato enriquecido (listados, tablas, enlaces a las vistas correspondientes de la plataforma).",
        "Mapa interactivo embebido en la propia conversación, que permite visualizar en tiempo real la posición de los conductores y los viajes consultados al asistente sin salir del chat.",
        "Mensajes de progreso personalizados en español durante la ejecución de cada operación interna del asistente, mostrando al usuario una descripción comprensible de la acción en curso en lugar de etiquetas técnicas.",
        "Opción de iniciar una nueva conversación, que limpia el contexto, los mensajes anteriores y los errores acumulados.",
        "Validación automática de los parámetros entregados al asistente antes de ejecutar cada operación, descartando valores vacíos o inválidos para evitar errores en la plataforma.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Persistencia del historial de consultas y respuestas por usuario en la base de datos.",
    ]),

    # ─────────────────── 15. Chat Viajes ───────────────────
    ("H1", "15. Chat Operativo de Viajes"),
    ("INTRO", "Permite la comunicación en tiempo real entre el conductor del viaje y el centro de control durante la ejecución del servicio."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Chat embebido en cada viaje, accesible tanto desde el portal del conductor como desde el módulo de administración de transporte.",
        "Mensajería en tiempo real con indicadores de mensaje enviado y mensaje leído.",
        "Persistencia de la conversación completa asociada al viaje para fines de auditoría.",
        "Soporte para envío de mensajes de texto y, opcionalmente, fotografías adjuntas.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Persistencia automática de todos los mensajes en la base de datos, asociados al viaje y al usuario emisor.",
    ]),

    # ─────────────────── 16. Usuarios y Permisos ───────────────────
    ("H1", "16. Administración de Usuarios y Permisos"),
    ("INTRO", "Permite gestionar las cuentas de los usuarios de la plataforma de administración y los módulos a los que cada uno tiene acceso, asegurando el principio de mínimo privilegio."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Creación de usuarios administradores con generación automática de contraseña inicial y envío de credenciales.",
        "Asignación de módulos visibles y accesibles para cada usuario, con filtrado automático de la navegación según los módulos otorgados.",
        "Edición de los datos del usuario y de los módulos asignados en cualquier momento.",
        "Cambio de contraseña gestionado por el propio usuario o por el administrador, con confirmación explícita.",
        "Eliminación de usuarios con confirmación explícita y registro de la acción.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Formulario de creación de usuario con validación de unicidad del correo electrónico.",
        "Persistencia de los módulos asignados en los metadatos del usuario.",
        "Registro de la fecha de creación, última modificación y último acceso.",
    ]),

    # ─────────────────── 17. Reportería ───────────────────
    ("H1", "17. Reportería Operativa"),
    ("INTRO", "Permite generar reportes consolidados de la operación del evento para análisis, rendición de cuentas y toma de decisiones, con acceso diferenciado según el rol del usuario dentro de la plataforma."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Reportes de transporte: viajes ejecutados por conductor, vehículo, delegación, tipo de viaje y rango de fechas.",
        "Reportes de participantes: distribución por delegación, tipo de cliente y servicios contratados.",
        "Reportes de calendario: actividades realizadas, asistencia y cumplimiento de horarios.",
        "Reportes financieros: facturación por proveedor según tarifas aplicadas y viajes ejecutados.",
        "Reportes de incidencias y asistencia: volumen, tiempos de respuesta y resolución, categorización y agente responsable.",
        "Exportación de todos los reportes a formato XLSX y PDF para su posterior análisis o distribución.",
        "**Acceso para los miembros del Comité Organizador**: los usuarios del Comité disponen de una vista específica dentro de la plataforma desde la cual pueden consultar y descargar directamente los reportes correspondientes a su ámbito de responsabilidad (operación general, transporte, participantes, incidencias, indicadores de cumplimiento), sin requerir solicitudes manuales al equipo operativo.",
        "**Filtros personalizables** por fecha, delegación, disciplina, proveedor, tipo de servicio y estado, permitiendo al usuario generar el reporte ajustado a su necesidad puntual.",
        "**Indicadores ejecutivos (KPIs)** con los totales y porcentajes clave de la operación presentados de forma visual para la toma de decisiones rápida.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Generación automática de los reportes a partir de la información operativa registrada en cada módulo, sin necesidad de carga adicional.",
        "Control de acceso a los reportes según el rol del usuario, asegurando que cada perfil (administrador, operador, Comité Organizador, proveedor) visualice únicamente la información que le corresponde.",
        "Registro de trazabilidad de cada descarga realizada (usuario, fecha, tipo de reporte) para fines de auditoría.",
    ]),

    # ─────────────────── 18. Auditoría ───────────────────
    ("H1", "18. Auditoría y Trazabilidad"),
    ("INTRO", "Permite mantener el registro histórico de todas las acciones relevantes ejecutadas en la plataforma, asegurando la trazabilidad de los cambios y la rendición de cuentas."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Bitácora por viaje con el detalle de todas las acciones ejecutadas (creación, asignación de conductor, cambio de vehículo, inicio de ruta, recogida, finalización, cancelación, observaciones manuales).",
        "Registro del usuario, fecha y hora de cada acción ejecutada.",
        "Visualización de la bitácora desde el módulo de transporte en formato de línea de tiempo.",
        "Persistencia indefinida de la bitácora para fines de auditoría.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Generación automática de las entradas de bitácora en respuesta a cada cambio relevante en los módulos operativos.",
        "Posibilidad de añadir entradas manuales desde el módulo de administración por parte del usuario autorizado.",
    ]),

    # ─────────────────── 19. Trazabilidad QR ───────────────────
    ("H1", "19. Trazabilidad del Deportista mediante Código QR"),
    ("INTRO", "Permite realizar el control de ingreso y la trazabilidad de los participantes del evento en cada uno de los puntos operativos (sedes deportivas, hoteles, zonas de alimentación, zonas restringidas) mediante el escaneo del código QR de la credencial digital portada en la aplicación móvil."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Cada participante del evento cuenta con una **credencial digital única** con código QR, accesible desde su aplicación móvil nativa.",
        "El personal autorizado en terreno (seguridad, acreditación, control de acceso) escanea el código QR del participante al momento de su ingreso a un punto operativo, mediante un dispositivo habilitado (teléfono, tablet o lector dedicado).",
        "Validación en tiempo real del código contra la base de datos central, con respuesta inmediata sobre la autorización del participante para ingresar al punto escaneado (autorizado, no autorizado, credencial vencida, participante no registrado).",
        "**Registro de cada escaneo** con fecha, hora, ubicación, punto de control y usuario que ejecutó la acción, generando una trazabilidad completa del recorrido del participante durante el evento.",
        "**Visualización del historial de ingresos** de cada participante desde el módulo de administración, disponible para los equipos de seguridad, operaciones y comité organizador.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Persistencia automática en la base de datos de cada escaneo realizado, con todos los atributos de contexto asociados.",
        "Generación de reportes consolidados de ingresos por punto de control, por delegación, por rango de fechas y por tipo de participante.",
        "Exportación de los reportes de trazabilidad a formato XLSX y PDF para fines de auditoría y rendición de cuentas.",
    ]),

    # ─────────────────── 20. Premiaciones ───────────────────
    ("H1", "20. Gestión de Premiaciones"),
    ("INTRO", "Permite administrar de forma centralizada las ceremonias de premiación asociadas a cada prueba deportiva del evento, definiendo el horario y ubicación de la ceremonia y asignando el equipo de participantes VIP encargado de entregar las medallas o reconocimientos a los deportistas premiados."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "**Asignación de fecha, hora y ubicación** de la ceremonia de premiación para cada prueba del calendario deportivo, de forma independiente al horario de la competencia.",
        "**Asignación del equipo de premiadores VIP**: desde el módulo de administración se selecciona a uno o más participantes registrados como tipo de cliente VIP para oficiar como entregadores del premio en cada ceremonia.",
        "Posibilidad de asignar **roles específicos dentro del equipo de premiación** (ej. entregador de medalla de oro, entregador de medalla de plata, entregador de medalla de bronce, autoridad oficial).",
        "**Visualización consolidada** del calendario de premiaciones del evento con filtros por fecha, disciplina, sede y estado.",
        "**Notificación automática a los VIP asignados**: los participantes VIP que integren un equipo de premiación reciben en su aplicación móvil la notificación con la información completa de la ceremonia (prueba, disciplina, fecha, hora, sede, ubicación exacta dentro de la sede, rol asignado, instrucciones especiales).",
        "**Integración con todos los portales de participantes**: la sección de premiaciones se encuentra disponible no solo en el portal VIP sino también en los portales del jefe de delegación y del atleta, con vista de calendario y vista de lista intercambiables, filtros por estado, disciplina, sede y búsqueda libre, agrupación por día y resumen de entregadores por rol.",
        "**Filtros por asistencia** en el portal VIP: confirmadas, pendientes y declinadas, con diferenciación visual por estado.",
        "**Confirmación de asistencia** por parte del VIP premiador desde su aplicación móvil, permitiendo al equipo operativo reaccionar ante ausencias o reasignaciones.",
        "**Registro histórico de las ceremonias ejecutadas** con los participantes VIP que efectivamente oficiaron el acto.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Formulario de creación de premiación asociado a una prueba del calendario deportivo, con validación de disponibilidad de sede y horario.",
        "Carga masiva de premiaciones mediante archivo XLSX con validación de fechas, sedes y pruebas.",
        "Edición individual y masiva de premiaciones con notificación automática a los VIP afectados ante cualquier cambio.",
        "Eliminación de premiaciones con confirmación explícita y trazabilidad de la acción.",
    ]),

    # ─────────────────── 21. Asistencia / Incidencias ───────────────────
    ("H1", "21. Centro de Asistencia y Gestión de Incidencias"),
    ("INTRO", "Permite canalizar, atender y resolver las incidencias, consultas y solicitudes de soporte levantadas por los conductores y los participantes del evento desde sus respectivos portales móviles, asegurando un flujo bidireccional de comunicación entre el usuario en terreno y el equipo operativo de la plataforma."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "**Apertura de salas de asistencia desde los portales móviles**: tanto el conductor como el participante (VIP, jefe de delegación y atleta) pueden abrir, desde su aplicación nativa, una sala de chat directa con el equipo operativo de Seven Arena para levantar incidencias (consultas, objetos perdidos, problemas operativos, solicitudes de apoyo, alertas de emergencia).",
        "**Bandeja centralizada de asistencia** en la plataforma web operativa: todas las salas abiertas por los usuarios aparecen en una bandeja consolidada dentro del módulo de administración, visible para los agentes operativos habilitados.",
        "**Asignación de agentes**: cada sala de asistencia puede ser tomada por un agente operativo, quien queda como responsable de la atención y resolución de la incidencia.",
        "**Diferenciación visual** de las salas según el origen (conductor o participante), la categoría de la incidencia y el nivel de prioridad asignado.",
        "**Chat en tiempo real** bidireccional entre el agente y el usuario, con soporte para envío de mensajes de texto, fotografías y archivos adjuntos como evidencia.",
        "**Categorización de incidencias**: las salas pueden clasificarse por tipo (consulta general, objeto perdido, incidencia operativa, crisis o emergencia, soporte al conductor, soporte al participante) para facilitar el análisis posterior y la reportería.",
        "**Estados de la incidencia**: abierta, en atención, escalada, resuelta, cerrada; con transiciones registradas y visibles en la ficha de la incidencia.",
        "**Escalamiento**: cuando una incidencia supera la capacidad de resolución del agente de primera línea, puede ser escalada a un coordinador o a una ruta de emergencia definida, con notificación automática al responsable correspondiente.",
        "**Historial completo de la conversación** persistido en la base de datos para fines de auditoría y análisis posterior.",
        "**Reportería de incidencias**: estadísticas de volumen, tiempo promedio de respuesta, tiempo promedio de resolución, incidencias por tipo, por usuario origen y por agente atendiente, exportable a formato XLSX y PDF.",
        "**Notificaciones en tiempo real**: los agentes reciben notificaciones dentro de la plataforma web cuando se abre una nueva sala o cuando llega un mensaje en una sala asignada; el usuario en el portal móvil recibe notificaciones push cuando el agente responde.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Persistencia automática en la base de datos de cada sala de asistencia creada, con sus mensajes, archivos adjuntos, categoría, agente asignado y trazabilidad completa de los cambios de estado.",
        "Registro del agente que tomó la incidencia, fecha y hora de apertura, fecha y hora de primera respuesta, fecha y hora de resolución.",
        "Posibilidad de añadir notas internas por parte del agente, visibles únicamente para el equipo operativo y no para el usuario origen.",
        "Cierre formal de la incidencia con registro del resultado (resuelta, derivada, no procedente) y observaciones finales.",
    ]),

    # ════════════════════════════════════════════════════════════════════════
    # NUEVOS MÓDULOS — implementados después de la v9 del documento
    # ════════════════════════════════════════════════════════════════════════

    # ─────────────────── 22. Cupones (NUEVO) ───────────────────
    ("H1", "22. Sistema de Cupones y Beneficios"),
    ("INTRO", "Permite ofrecer a los participantes del evento un catálogo digital de beneficios y descuentos en comercios aliados, gestionando todo el flujo desde la publicación del cupón hasta su canje en el local del comercio."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Catálogo de cupones con fotografía, comercio asociado, categoría (comida, entretenimiento, tienda, otros), tipo de descuento (porcentaje, monto fijo, gratuito), términos y condiciones, vigencia (fecha de inicio y término) y límite máximo de canjes por participante.",
        "Reclamo de cupón desde el portal del participante (VIP, jefe de delegación o atleta) con generación inmediata de un código único y un código QR asociado al usuario.",
        "Visualización del cupón reclamado en pantalla completa dentro del portal, con el código QR escaneable por el comercio, código de respaldo destacado, instrucciones de canje y términos del beneficio.",
        "Estados del cupón reclamado: reclamado, canjeado, expirado y anulado, visibles para el participante en la sección Mis Cupones de su portal.",
        "Portal independiente para el comercio aliado, que permite escanear el código QR del cupón, validar su autenticidad contra la plataforma y registrar el canje efectivo.",
        "Acceso sin inicio de sesión adicional: el catálogo y el reclamo reutilizan la sesión del portal del participante, sin requerir credenciales adicionales.",
        "Disponibilidad transversal para todos los tipos de cliente (atleta, jefe de delegación, VIP) sin filtros por rol.",
        "Administración del catálogo desde el módulo central: alta, edición, baja, control de stock y vigencia.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Formulario de alta de cupón con fotografía, datos del comercio aliado, categoría, tipo y valor de descuento, vigencia y límite por participante.",
        "Persistencia automática de cada cupón reclamado con el participante, fecha, código único, código QR, fecha de expiración y estado actual.",
        "Registro de cada canje con fecha, hora y comercio que ejecutó la operación, para fines de auditoría y conciliación.",
    ]),

    # ─────────────────── 23. Workforce (NUEVO) ───────────────────
    ("H1", "23. Gestión de Personal Operativo y Voluntarios"),
    ("INTRO", "Permite administrar al personal operativo del evento que no es participante deportivo (personal de servicio, voluntarios, agentes de soporte, supervisores) y controlar el inventario de productos y materiales entregados a cada persona durante la operación."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Registro del personal operativo con categoría (personal de servicio, voluntario, supervisor, agente), fotografía, datos de contacto, área asignada y estado de validación.",
        "Catálogo de productos asignables (uniformes, kits, credenciales físicas, materiales operativos) con código de barras único por producto.",
        "Registro de entregas asociando producto y persona, con fecha, responsable de la entrega y observaciones.",
        "Visualización del personal y de los productos con tablas operativas que incluyen avatar con iniciales, indicadores de progreso, etiquetas de categoría, datos de contacto y estado de validación.",
        "Impresión de etiquetas con código de barras en formato individual o por lote, en hoja imprimible lista para enviar al impresor.",
        "Reportería de entregas, faltantes y estado por persona y por producto.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Formulario de alta del personal operativo con validación de duplicados.",
        "Catálogo de productos con código de barras autogenerado o ingresado manualmente.",
        "Registro de entrega con trazabilidad completa: identificación del responsable, persona receptora, fecha y producto entregado.",
    ]),

    # ─────────────────── 24. Acreditaciones (NUEVO) ───────────────────
    ("H1", "24. Gestión Administrativa de Acreditaciones"),
    ("INTRO", "Complementa al punto 19 (trazabilidad mediante código QR en terreno) con la gestión administrativa previa a la operación: solicitud, aprobación y generación de la credencial digital del participante o conductor."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Solicitud de acreditación para participantes y conductores, con asignación de los tipos de acceso correspondientes (sedes deportivas, zonas restringidas, hoteles, áreas de prensa).",
        "Estados de la acreditación: pendiente, aprobada, rechazada y vencida, con registro de cada transición y del usuario que la ejecutó.",
        "Generación e impresión de la credencial digital con fotografía del titular, datos personales, código QR, zonas de acceso permitidas e identidad institucional del evento.",
        "Resolución automática de la fotografía del titular a partir de los datos del participante, con visualización de un marcador de ausencia de fotografía cuando el registro no la contiene.",
        "Previsualización de la credencial dentro de la propia plataforma, sin necesidad de abrir una ventana adicional, con opción de imprimirla desde la misma vista.",
        "Plantilla adaptable de la credencial: formato A4 apaisado para impresión y visualización optimizada en dispositivos móviles.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Formulario de solicitud con validación de unicidad por participante y tipos de acceso compatibles.",
        "Persistencia del histórico de acreditaciones por participante (vigentes, vencidas y revocadas).",
        "Integración con el módulo 19 (trazabilidad mediante código QR) para que cada escaneo en terreno valide la acreditación contra el estado actual.",
    ]),

    # ─────────────────── 25. Centro de Ayuda (NUEVO) ───────────────────
    ("H1", "25. Centro de Ayuda y Documentación"),
    ("INTRO", "Permite a los usuarios administradores consultar dentro de la plataforma la documentación operativa del evento, incluyendo glosario de términos, roles, regiones, recintos, hoteles, coordinadores y procedimientos generales."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Secciones temáticas filtrables: general, buses, áreas, glosario, roles, regiones, recintos, hoteles y coordinadores.",
        "Incorporación del Cuaderno de Cargo del evento como sección principal, con cerca de cien entradas estructuradas (término, detalles, etiquetas) y soporte trilingüe en español, inglés y portugués para las etiquetas de categorías.",
        "Búsqueda libre dentro del contenido por término, descripción o etiqueta.",
        "Visualización del documento original del Cuaderno de Cargo dentro de la plataforma como referencia descargable en formato PDF.",
        "Acceso desde el menú principal de la plataforma, disponible para todos los roles administrativos sin requerir permisos especiales.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Contenido administrado de forma estructurada y versionada, con validación de campos obligatorios para cada entrada.",
        "Posibilidad de ampliar las categorías y entradas en versiones posteriores de la plataforma sin requerir cambios en la base de datos.",
    ]),

    # ─────────────────── 26. Inicio Guiado (NUEVO) ───────────────────
    ("H1", "26. Inicio Guiado al Usuario Administrador"),
    ("INTRO", "Permite acompañar a los nuevos usuarios administradores en sus primeros pasos dentro de la plataforma mediante un recorrido interactivo personalizado por rol, recomendando los módulos relevantes según su perfil y los objetivos declarados."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Recorrido estructurado en seis pasos: bienvenida, selección del rol del usuario, selección de objetivos, recomendación de tareas, sugerencias y atajos, y pantalla de finalización.",
        "Seis roles configurados (Director, Coordinador de Operaciones, Coordinador de Transporte, Coordinador de Acreditación, Coordinador de Hospedaje y Coordinador de Alimentación), cada uno con su ícono representativo.",
        "Listado de objetivos seleccionables por el usuario, con etiquetas que indican el rol al que aplican.",
        "Listado de tareas recomendadas con enlaces directos a los módulos correspondientes de la plataforma.",
        "Persistencia del progreso en el navegador del usuario para permitir reanudar el recorrido en una sesión posterior.",
        "Animaciones y elementos visuales (partículas, transiciones, indicadores de paso activo, efecto de confeti al completar) que entregan una experiencia diferenciada respecto a un manual estático.",
        "Indicadores de paso clickeables que permiten navegar entre las etapas del recorrido.",
        "Estimación visual del tiempo restante para completar el recorrido.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Persistencia del estado del recorrido en el navegador del usuario para continuidad entre sesiones.",
        "Posibilidad de restablecer el recorrido manualmente desde la propia pantalla.",
        "Acceso para todos los usuarios administradores desde el menú principal de la plataforma.",
    ]),

    # ─────────────────── 27. Operatividad Diaria (NUEVO) ───────────────────
    ("H1", "27. Operatividad Diaria del Transporte"),
    ("INTRO", "Permite consolidar en una sola pantalla la planificación, asignación y seguimiento de los viajes del día, entregando al equipo operativo las herramientas necesarias para administrar la jornada de transporte de manera eficiente."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Vista consolidada del día con todos los viajes en estados activos (solicitado, programado, en ruta y recogido) y sus indicadores clave.",
        "Asignación automática de conductores con criterios operativos: capacidad del vehículo, disponibilidad horaria y carga acumulada de cada conductor en la jornada.",
        "Descarga de una plantilla en formato XLSX con todos los campos necesarios para la carga masiva del transporte del día, incluyendo filas de ejemplo y campos vacíos para los datos que completa el sistema (conductor, teléfono, patente).",
        "Filtrado del listado por estado, conductor, vehículo, ruta y franja horaria.",
        "Edición masiva sobre selección múltiple para asignar o reasignar conductores en lote.",
        "Indicadores agregados con el total de viajes del día, viajes asignados, viajes sin asignar y porcentaje de cobertura.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Carga masiva mediante archivo XLSX usando la plantilla descargable, con validación previa de cada fila y reporte de errores antes de la confirmación.",
        "Persistencia de la asignación automática en la bitácora de cada viaje, identificando que la operación fue ejecutada por la plataforma.",
        "Posibilidad de reasignación manual en cualquier momento previo al inicio del viaje.",
    ]),

    # ─────────────────── 28. Acciones de SofIA (NUEVO) ───────────────────
    ("H1", "28. Registro de Acciones del Asistente de Inteligencia Artificial"),
    ("INTRO", "Permite auditar de forma completa cada operación que el asistente de inteligencia artificial (SofIA) ejecuta sobre la plataforma, garantizando la trazabilidad, el control y la rendición de cuentas sobre el accionar de la inteligencia artificial."),
    ("H2", "a) Funcionalidad"),
    ("UL", [
        "Registro histórico de cada operación ejecutada por el asistente: viajes creados, modificados o cancelados, hoteles asignados, premiaciones programadas, cupones publicados, notificaciones enviadas, entre otras.",
        "Filtrado del listado por tipo de operación, por fecha, por usuario que originó la consulta y por estado de la ejecución (exitosa o con error).",
        "Detalle expandible de cada operación con los parámetros enviados, la respuesta devuelta por la plataforma y la consulta original que la originó.",
        "Indicadores agregados del uso del asistente: operaciones por día, operaciones por tipo, porcentaje de éxito y funciones más utilizadas.",
        "Exportación del registro a formato XLSX para su posterior análisis o reportería ejecutiva.",
    ]),
    ("H2", "b) Módulo de registro"),
    ("UL", [
        "Persistencia automática en la base de datos de cada operación ejecutada por el asistente, con los parámetros validados y la respuesta completa de la plataforma.",
        "Asociación de cada operación con la consulta de origen para facilitar la auditoría contextual del comportamiento del asistente.",
    ]),

    # ─────────────────── 29. Arquitectura (era 22) ───────────────────
    ("H1", "29. Arquitectura Tecnológica (Descripción General)"),
    ("INTRO", "La plataforma Seven Arena ha sido construida sobre una arquitectura moderna de tres capas (frontend, backend y base de datos), desplegada en infraestructura cloud de alta disponibilidad y diseñada para soportar la operación simultánea de grandes volúmenes de usuarios durante el evento deportivo."),
    ("H2", "a) Backend y servicios"),
    ("UL", [
        "Capa de servicios desarrollada con tecnologías modernas basadas en Node.js y TypeScript, que aseguran una arquitectura modular, mantenible y escalable.",
        "Exposición de los servicios mediante una interfaz estándar de integración, con validación estricta de entrada y autenticación basada en tokens seguros.",
        "Transmisión de eventos en tiempo real (ubicaciones GPS, notificaciones, mensajería) para soportar la operación en campo.",
        "Integración con servicios externos para funcionalidades complementarias: mapas y geocodificación, seguimiento de vuelos, inteligencia artificial y almacenamiento de archivos.",
        "Despliegue en infraestructura cloud con escalamiento automático y monitoreo continuo.",
    ]),
    ("H2", "b) Frontend (Plataforma Web de Administración)"),
    ("UL", [
        "Plataforma de administración desarrollada con React y TypeScript, asegurando una experiencia de usuario fluida y alta mantenibilidad del código.",
        "Interfaz responsiva con sistema de componentes reutilizables e identidad visual coherente en todos los módulos.",
        "Despliegue con entrega de contenido mediante red de distribución global, asegurando baja latencia desde cualquier ubicación.",
    ]),
    ("H2", "c) Base de Datos"),
    ("UL", [
        "Motor PostgreSQL gestionado en infraestructura cloud, con respaldos automáticos diarios y replicación de alta disponibilidad.",
        "Modelo de datos relacional normalizado, organizado en esquemas funcionales por área (participantes y eventos, transporte, logística y alimentación, telemetría, configuración y autenticación).",
        "Control de acceso a la información según el rol del usuario, asegurando que cada participante y proveedor acceda únicamente a los datos que le corresponden.",
        "Optimización de consultas sobre los datos críticos de la operación (viajes, posiciones GPS, participantes, calendario) para asegurar tiempos de respuesta bajos incluso con volúmenes elevados.",
    ]),
    ("H2", "d) Seguridad"),
    ("UL", [
        "Autenticación centralizada con gestión de sesiones, recuperación de contraseña y soporte para mecanismos de doble factor.",
        "Cifrado de las comunicaciones mediante HTTPS/TLS en todos los puntos de acceso.",
        "Almacenamiento de archivos sensibles (documentos, fotografías) en repositorios privados con políticas de acceso por rol.",
        "Registro de auditoría de las acciones relevantes ejecutadas en la plataforma, incluyendo las operaciones realizadas por el asistente de inteligencia artificial (ver sección 28).",
        "Aplicación de buenas prácticas de seguridad en el desarrollo, conforme a estándares reconocidos de la industria (OWASP).",
    ]),
    ("H2", "e) Escalabilidad y Disponibilidad"),
    ("UL", [
        "Infraestructura cloud con escalamiento horizontal automático según la demanda.",
        "Caché de alto rendimiento para las consultas más frecuentes.",
        "Monitoreo continuo del estado de los servicios con alertas automáticas ante cualquier degradación.",
        "Capacidad de soportar grandes volúmenes de usuarios concurrentes durante los picos de operación del evento deportivo.",
    ]),
    ("H2", "f) Inteligencia Artificial integrada"),
    ("UL", [
        "La plataforma integra un asistente conversacional basado en modelos avanzados de Inteligencia Artificial (SofIA), que permite a los usuarios administradores consultar información operativa en lenguaje natural y ejecutar acciones controladas.",
        "El asistente accede de forma controlada a los datos operativos de la plataforma mediante un conjunto de herramientas internas definidas y seguras, asegurando que las respuestas reflejen siempre el estado actualizado del sistema.",
        "Mapa interactivo embebido en la conversación con el asistente, para visualizar la posición de los conductores y los viajes consultados sin salir del chat.",
        "Mensajes de progreso personalizados en español durante la ejecución de cada operación del asistente, en lugar de etiquetas técnicas.",
        "Capacidad de ampliar las funcionalidades del asistente mediante la incorporación de nuevas herramientas y flujos conversacionales según las necesidades operativas del evento.",
        "Aplicación de buenas prácticas de seguridad y privacidad en el manejo de la información enviada al modelo de lenguaje, cumpliendo con los estándares de protección de datos.",
        "Auditoría completa de las operaciones ejecutadas por el asistente, descrita en la sección 28.",
    ]),
    ("H2", "g) Soporte multilenguaje"),
    ("UL", [
        "La plataforma ha sido diseñada con **soporte multilenguaje nativo**, permitiendo la presentación de la interfaz y de los contenidos operativos en distintos idiomas según las necesidades de cada usuario.",
        "Idiomas soportados inicialmente: **español, inglés y portugués**, con posibilidad de incorporar idiomas adicionales de las delegaciones participantes en el evento deportivo.",
        "Selección del idioma por parte de cada usuario desde su perfil, con persistencia de la preferencia entre sesiones.",
        "Traducción automática de las notificaciones, mensajes operativos y reportes según el idioma preferido del destinatario.",
        "Aplicación del soporte multilenguaje tanto en la plataforma web de administración como en las aplicaciones móviles nativas de los cuatro portales de usuario.",
    ]),

    # ─────────────────── 30. App Móvil (era 23) ───────────────────
    ("H1", "30. Aplicación Móvil Nativa (Portales de Usuario)"),
    ("INTRO", "Como complemento a la plataforma web de administración, los cuatro portales de usuario final (Conductor, VIP, Jefe de Delegación y Atleta) serán desarrollados como aplicaciones móviles nativas para los sistemas operativos iOS y Android. Esta migración responde a la necesidad operativa de asegurar la continuidad del tracking GPS en segundo plano, la recepción inmediata de notificaciones push y la entrega de una experiencia de usuario acorde a los estándares de las aplicaciones móviles modernas, condiciones que no es posible garantizar de forma confiable mediante una aplicación web."),
    ("H2", "a) Arquitectura general"),
    ("UL", [
        "Las aplicaciones móviles nativas se integran directamente con la plataforma central Seven Arena, sin reemplazar el sistema de administración web actual utilizado por el equipo operativo.",
        "Toda la lógica de negocio permanece centralizada en la plataforma; las aplicaciones móviles actúan como clientes que consumen los servicios existentes, asegurando consistencia de la información en tiempo real.",
    ]),
    ("H2", "b) Stack tecnológico propuesto"),
    ("UL", [
        "Desarrollo multiplataforma con **React Native**, permitiendo una única base de código para los sistemas operativos **iOS** y **Android** con alto porcentaje de reutilización.",
        "**TypeScript** como lenguaje principal, consistente con el resto de la plataforma.",
        "Publicación en las tiendas oficiales **App Store** (Apple) y **Google Play** (Google).",
    ]),
    ("H2", "c) Funcionalidades nativas del dispositivo"),
    ("UL", [
        "**Geolocalización en segundo plano**, manteniendo el envío continuo de la posición del conductor al centro de control incluso con la aplicación minimizada o la pantalla apagada.",
        "**Notificaciones push** con apertura directa en la pantalla adecuada al tocar la notificación.",
        "**Cámara y galería** del dispositivo para la carga de fotografías de perfil, del vehículo, de documentos y de evidencia de jornada laboral.",
        "**Mapas interactivos** con marcadores, rutas y posición en tiempo real.",
        "**Biometría opcional** (Face ID / Touch ID) como mejora de seguridad para el inicio de sesión.",
        "**Credencial digital con código QR** para la acreditación de ingreso a sedes, hoteles y servicios de alimentación.",
    ]),
    ("H2", "d) Portal Conductor — App Nativa"),
    ("UL", [
        "**Autenticación por código único**: el conductor ingresa un código de 6 caracteres (los últimos 6 caracteres de su identificador de participante del proveedor), distinto del mecanismo estándar de email y contraseña.",
        "Bandeja de viajes asignados filtrable por estado: Hoy, En curso, Programados e Historial, con marcado visual de viajes no vistos (estilo bandeja de correo).",
        "Flujo operativo completo del viaje: iniciar ruta, validación del pasajero mediante **código de pickup**, marcar pasajero recogido, marcar dejado en destino y completar viaje.",
        "**Envío automático de ubicación GPS cada 5 segundos** mientras el viaje se encuentra en estado *En ruta* o *Recogido*, funcionando con la aplicación minimizada o la pantalla apagada.",
        "Gestión del perfil del conductor: carga de **foto de perfil, 12 documentos personales y del vehículo** (cédula, licencia de conducir, antecedentes, SOAP, permiso de circulación, padrón, foto del vehículo, contrato entre proveedores, entre otros), y **2 fotos diarias obligatorias** para monitorear la jornada laboral.",
        "**Credencial digital con código QR** para la acreditación del conductor.",
        "Soporte para **viajes tipo Disposición 12 horas**, con ejecución paralela de viajes tipo Transfer In/Out dentro del periodo de disposición y cálculo diferenciado del costo final según los servicios efectivamente realizados.",
        "**Chat en tiempo real** con el pasajero durante los viajes activos y **sala de asistencia al conductor** atendida por un agente operativo con soporte para envío de mensajes, archivos y fotografías como evidencia.",
        "**Historial de viajes** con filtros por fecha, tipo y calificación recibida, además de estadísticas personales (total de viajes y calificación promedio).",
        "Notificaciones push ante: nueva asignación de viaje, nuevo mensaje del pasajero o modificación/cancelación de viajes programados.",
    ]),
    ("H2", "e) Portal Solicitud de Vehículo (VIP) — App Nativa"),
    ("UL", [
        "Autenticación mediante **Supabase Auth** con email y contraseña.",
        "Pantalla principal orientada al **formulario de solicitud de vehículo**, con selección de tipo (Sedán 4, SUV 6, Van 10/15/19, Minibús 33, Bus 64), validación de capacidad, dirección de origen con **autocompletado Google Places**, sede destino, fecha, hora, observaciones y toggle de *ida y vuelta* con retorno a destino distinto.",
        "**Modificación o cancelación** de solicitudes con máximo de anticipación definido por la regla de negocio.",
        "**Seguimiento en tiempo real** del viaje activo con mapa, posición del conductor, datos del vehículo asignado (nombre del conductor, foto, patente, modelo) y tiempo estimado de llegada.",
        "**Chat en tiempo real** con el conductor durante los viajes activos, con notificaciones push de nuevos mensajes.",
        "**Calificación del viaje** al finalizar (1 a 5 estrellas y comentario opcional).",
        "Secciones complementarias: **Actividades** (viajes en curso, programados e historial), **Sedes** (directorio con fotos y mapas), **Hoteles** (Villa Panamericana y alojamientos), **Comida** (menús y lugares, sin filtros por tipo de cliente), **Calendario** deportivo, **Premiaciones** (calendario + lista intercambiables con CTAs de confirmar/declinar asistencia), **Cupones** (catálogo y Mis cupones con QR), y **Cuenta** con credencial digital QR inline.",
        "**Sala de asistencia al participante** para levantar incidencias, consultas u objetos perdidos con un agente de operaciones.",
    ]),
    ("H2", "f) Portal Jefe de Delegación — App Nativa"),
    ("UL", [
        "Autenticación mediante **Supabase Auth** con email y contraseña.",
        "**Pestañas operativas**: Itinerario, Actividades, Calendario, Sedes, Alimentación, Premiaciones, Cupones, Delegación y Cuenta.",
        "**Itinerario personal** con vista consolidada del día, ordenado cronológicamente, incluyendo viajes, actividades deportivas, comidas y eventos programados.",
        "**Gestión de actividades**: creación y seguimiento de viajes para los miembros de la delegación, con mapa en tiempo real, chat con el conductor y bitácora de cambios en el historial.",
        "**Calendario deportivo** con vista mensual, detalle diario por disciplina, pruebas, sedes y horarios, con filtros por disciplina o delegación.",
        "**Directorio de sedes** con fotografías, direcciones, ubicación en mapa e información de contacto.",
        "**Menús de alimentación** por fecha y tipo (desayuno, almuerzo, cena, coffee break), con lugares de servicio, ubicación y capacidad. Sin filtros por tipo de cliente.",
        "**Pestaña Premiaciones** con vista calendario + lista y filtros completos.",
        "**Pestaña Cupones** con catálogo y reclamo de QR.",
        "**Administración de la delegación**: listado completo de miembros, información individual (nombre, disciplina, rol, estado de acreditación) y datos agregados.",
        "**Cuenta y credencial digital QR** con foto de perfil y datos personales, accesible dentro de la propia aplicación sin necesidad de abrir una ventana adicional.",
        "**Sala de asistencia** para levantar incidencias con un agente de operaciones.",
        "Notificaciones push ante: confirmación de solicitud de viaje, asignación de conductor, estados del viaje, mensajes del conductor, cambios en el calendario y alertas de acreditación.",
    ]),
    ("H2", "g) Portal Atleta — App Nativa"),
    ("UL", [
        "Autenticación mediante **Supabase Auth** con email y contraseña.",
        "**Pestañas operativas**: Actividades, Calendario, Sedes, Alimentación, Premiaciones, Cupones y Cuenta.",
        "**Actividades**: consulta de viajes propios como pasajero (cuando el jefe de delegación lo ha vinculado), seguimiento en tiempo real, chat con el conductor y calificación al finalizar.",
        "**Calendario deportivo** con vista mensual y énfasis en las actividades donde el atleta participa.",
        "**Directorio de sedes** con fotos, direcciones y mapas.",
        "**Menús de alimentación** diarios sin filtros por tipo de cliente.",
        "**Pestaña Premiaciones** con vista calendario + lista.",
        "**Pestaña Cupones** con catálogo y reclamo de QR.",
        "**Credencial digital con código QR** para acreditación de ingreso a sedes, hoteles y alimentación, accesible dentro de la propia aplicación sin necesidad de abrir una ventana adicional.",
        "**Sala de asistencia al participante** para incidencias y consultas.",
        "Notificaciones push ante: viaje asignado donde el atleta es pasajero, estados del viaje, mensajes del conductor, cambios en el calendario y actualizaciones de acreditación.",
    ]),
    ("H2", "h) Autenticación y resolución del portal"),
    ("UL", [
        "**Portal Conductor**: autenticación mediante un código único de 6 caracteres entregado al conductor al momento de su registro, con persistencia de la sesión para facilitar el acceso en siguientes aperturas.",
        "**Portales VIP, Jefe de Delegación y Atleta**: autenticación mediante email y contraseña, reutilizando las credenciales ya asignadas al participante en la plataforma central.",
        "**Resolución automática del portal** tras el inicio de sesión, según los atributos del participante:",
        "— Participante asociado a proveedor de transporte con **esChofer = sí** → Portal Conductor.",
        "— Participante con **tipoCliente = VIP** → Portal Solicitud de Vehículo.",
        "— Participante con **tipoCliente = TA** y **jefeDelegación = sí** → Portal Jefe de Delegación.",
        "— Participante con **tipoCliente = TA** y **jefeDelegación = no** → Portal Atleta.",
    ]),
    ("H2", "i) Requisitos funcionales comunes"),
    ("UL", [
        "**Sesión persistente** entre cierres de la aplicación, con cierre manual y opción de biometría (Face ID / Touch ID).",
        "**Notificaciones push** con apertura en la pantalla adecuada al tocar la notificación (*deep linking*).",
        "**Solicitud amigable de permisos** de ubicación, cámara y notificaciones.",
        "**Funcionamiento con red inestable** mediante reintentos automáticos e indicación visual del estado de conexión.",
        "**Compresión automática de imágenes** antes de subirlas al servidor.",
        "**Interfaz en español** (con soporte multilenguaje), diseño consistente con la identidad de marca, soporte para modo horizontal y vertical, y accesibilidad básica.",
    ]),
    ("H2", "j) Reglas de negocio aplicables"),
    ("UL", [
        "Un viaje se crea en estado **Solicitado**, pasa a **Programado** al asignar conductor y vehículo, y avanza secuencialmente a **En ruta**, **Recogido**, **Dejado en hotel** y **Completado**.",
        "El usuario puede **modificar o cancelar un viaje únicamente si faltan más de 2 horas** para su hora programada.",
        "Cada modificación, cancelación o asignación se registra en la **bitácora del viaje**.",
        "Un viaje de **ida y vuelta** genera dos viajes enlazados (padre e hijo).",
        "Al asignar conductor, solo se muestran conductores cuyo vehículo tenga **capacidad suficiente** para los pasajeros solicitados.",
        "El **chat solo está disponible** durante los estados *En ruta* y *Recogido*.",
        "La **calificación del conductor** (1 a 5 estrellas y comentario opcional) se realiza al finalizar el viaje y alimenta el ranking operacional.",
    ]),
    ("H2", "k) Continuidad operativa con la plataforma central"),
    ("UL", [
        "Las aplicaciones móviles se integran de forma transparente con la plataforma Seven Arena existente, reutilizando el modelo de datos, la lógica de negocio y los mecanismos de seguridad ya operativos.",
        "La información registrada desde las aplicaciones móviles (viajes, ubicaciones, mensajes, documentos, incidencias) se refleja de forma inmediata en la plataforma web operativa utilizada por el equipo de administración.",
        "La infraestructura cloud actual soporta el tráfico adicional de las aplicaciones móviles sin requerir modificaciones mayores, asegurando la continuidad del servicio durante la operación del evento.",
    ]),
    ("H2", "l) Publicación y gestión de versiones"),
    ("UL", [
        "Publicación en **App Store** (Apple) y **Google Play** (Google) bajo las cuentas de desarrollador de la organización.",
        "Gestión de **versiones y actualizaciones automáticas** mediante los mecanismos nativos de cada tienda.",
        "Ambiente de **pruebas internas** (TestFlight para iOS, Internal Testing para Android) previo a la publicación pública de cada versión.",
        "**Documentación técnica y funcional completa** del proyecto disponible para el proveedor a cargo del desarrollo (referencia: documento *Seven Arena — Documentación Técnica y Funcional de Portales Móviles*).",
    ]),
]


# ──────────────────────────────────────────────────────────────────────────────
# CONSTRUCCIÓN DEL DOCX
# ──────────────────────────────────────────────────────────────────────────────

def set_paragraph_spacing(p, before=0, after=4):
    """Configura espaciado antes/después en puntos."""
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def add_runs_with_bold(p, text):
    """Procesa **negritas** dentro del texto y agrega runs adecuados."""
    parts = text.split("**")
    for i, chunk in enumerate(parts):
        if not chunk:
            continue
        run = p.add_run(chunk)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True


def set_list_bullet(p):
    """Asigna formato de viñeta nativa a un párrafo."""
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "1")
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)


def build(doc):
    # Estilos base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    for kind, payload in DOC:
        if kind == "TITLE":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(payload)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            set_paragraph_spacing(p, before=0, after=4)

        elif kind == "SUBTITLE":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(payload)
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            set_paragraph_spacing(p, before=0, after=18)

        elif kind == "H1":
            p = doc.add_paragraph()
            run = p.add_run(payload)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x0F, 0x9E, 0x87)  # teal-ish
            set_paragraph_spacing(p, before=14, after=4)

        elif kind == "INTRO":
            p = doc.add_paragraph()
            add_runs_with_bold(p, payload)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            set_paragraph_spacing(p, before=2, after=6)

        elif kind == "H2":
            p = doc.add_paragraph()
            run = p.add_run(payload)
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            set_paragraph_spacing(p, before=6, after=2)

        elif kind == "P":
            p = doc.add_paragraph()
            add_runs_with_bold(p, payload)
            set_paragraph_spacing(p, before=2, after=4)

        elif kind == "UL":
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.6)
                add_runs_with_bold(p, item)
                set_paragraph_spacing(p, before=0, after=2)


if __name__ == "__main__":
    import os
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "FUNCIONALIDADES_PLATAFORMA_v10.docx")
    doc = Document()
    build(doc)
    doc.save(out_path)
    print(f"OK: {out_path}")
