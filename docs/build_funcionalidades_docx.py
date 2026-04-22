"""Convert FUNCIONALIDADES_PLATAFORMA.md to DOCX in formal tender style."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
MD = (HERE / "FUNCIONALIDADES_PLATAFORMA.md").read_text(encoding="utf-8")
OUT = HERE / "FUNCIONALIDADES_PLATAFORMA_v9.docx"

doc = Document()

for section in doc.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

BLACK = RGBColor(0x00, 0x00, 0x00)


def set_underline(run):
    run.underline = True


def add_module_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK


def add_subsection_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def add_paragraph_text(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(p, text)


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(p, text)


def add_inline(paragraph, text):
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# Cover title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("Funcionalidades de la Plataforma Seven Arena")
trun.bold = True
trun.font.size = Pt(16)
doc.add_paragraph()

lines = MD.split("\n")
i = 0
# Skip the H1 since we already added the cover
while i < len(lines) and not lines[i].startswith("## "):
    i += 1

while i < len(lines):
    line = lines[i].rstrip()

    if line.startswith("## "):
        add_module_title(line[3:].strip())
        i += 1
        continue

    if line.startswith("### "):
        add_subsection_title(line[4:].strip())
        i += 1
        continue

    if line.startswith("- "):
        add_bullet(line[2:].strip())
        i += 1
        continue

    if line.strip() == "---" or not line.strip():
        i += 1
        continue

    add_paragraph_text(line.strip())
    i += 1

doc.save(OUT)
print(f"Saved: {OUT}")
